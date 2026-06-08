from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "prd_agent_mvp_operation_guide.md"
OUTPUT = ROOT / "docs" / "PRD生成Agent_MVP操作步骤与文档评审.docx"


def set_east_asia_font(run, font_name="Microsoft YaHei"):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_shading(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def add_bottom_border(paragraph, color="B7C6D8"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    ppr.append(pbdr)


def style_run(run, size=11, bold=False, color="000000", mono=False):
    set_east_asia_font(run, "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    if mono:
        run.font.name = "Consolas"
        rpr = run._element.get_or_add_rPr()
        rpr.rFonts.set(qn("w:ascii"), "Consolas")
        rpr.rFonts.set(qn("w:hAnsi"), "Consolas")


def add_inline_runs(paragraph, text, size=11, bold=False, color="000000"):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            style_run(run, size=size, color=color, mono=True)
        else:
            run = paragraph.add_run(part)
            style_run(run, size=size, bold=bold, color=color)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(header, after=2)
    run = header.add_run("PRD 生成 Agent MVP 操作步骤与文档评审")
    style_run(run, size=9, color="5B677A")
    add_bottom_border(header, "D5DCE6")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("MVP v0.1 验证文档")
    style_run(run, size=9, color="5B677A")


def add_title_block(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(title, before=0, after=6)
    run = title.add_run("PRD 生成 Agent")
    style_run(run, size=24, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, before=0, after=12)
    run = subtitle.add_run("MVP 操作步骤与需求文档评审")
    style_run(run, size=14, color="1F4D78")

    meta = doc.add_paragraph()
    set_paragraph_spacing(meta, after=16)
    add_shading(meta, "F4F6F9")
    run = meta.add_run("用途：用于 2-3 周内验证 PRD 生成 Agent 的核心闭环，并补充原始需求文档的落地缺口。")
    style_run(run, size=10.5, color="263445")


def convert_markdown(md_text):
    doc = Document()
    configure_document(doc)
    add_header_footer(doc)
    add_title_block(doc)

    in_code = False
    code_lines = []
    for raw in md_text.splitlines():
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=2, after=8, line=1.0)
                add_shading(p, "F2F4F7")
                run = p.add_run("\n".join(code_lines))
                style_run(run, size=9.5, color="222222", mono=True)
                code_lines = []
                in_code = False
            else:
                in_code = True
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line:
            continue

        if line.startswith("# "):
            continue
        if line.startswith("## "):
            p = doc.add_paragraph(line[3:].strip(), style="Heading 1")
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(line[4:].strip(), style="Heading 2")
            continue

        number_match = re.match(r"^(\d+)\.\s+(.*)", line)
        if number_match:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, number_match.group(2), size=11)
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line[2:].strip(), size=11)
            continue

        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=6, line=1.1)
        add_inline_runs(p, line, size=11)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    convert_markdown(SOURCE.read_text(encoding="utf-8"))
    print(OUTPUT)
