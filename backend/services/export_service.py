from __future__ import annotations

import re
from pathlib import Path

from backend.models import SessionState


ROOT = Path(__file__).resolve().parents[2]
EXPORT_DIR = ROOT / "exports"
EXPORT_DIR.mkdir(exist_ok=True)

FONT_SIZE_PT = 9
CN_FONT = "SimSun"
CN_FONT_FALLBACK = "宋体"
EN_FONT = "Calibri"


def _safe_name(value: str) -> str:
    value = re.sub(r"[^\w\u4e00-\u9fff-]+", "_", value).strip("_")
    return value[:40] or "prd"


def _docx_set_run_font(run, bold: bool = False) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.bold = bold
    run.font.name = EN_FONT
    run.font.size = Pt(FONT_SIZE_PT)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), EN_FONT)
    rfonts.set(qn("w:hAnsi"), EN_FONT)
    rfonts.set(qn("w:eastAsia"), CN_FONT)
    rfonts.set(qn("w:cs"), EN_FONT)


def _docx_style_paragraph(paragraph, before: int = 0, after: int = 4) -> None:
    from docx.shared import Pt

    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.15


def _docx_configure_document(doc) -> None:
    from docx.enum.section import WD_SECTION_START
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = EN_FONT
        style.font.size = Pt(FONT_SIZE_PT)
        style._element.rPr.rFonts.set(qn("w:ascii"), EN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), EN_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), EN_FONT)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    for heading in ["Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[heading].font.bold = True


def _docx_add_paragraph(doc, text: str, style: str | None = None, bold: bool = False):
    paragraph = doc.add_paragraph(style=style)
    _docx_style_paragraph(paragraph)
    run = paragraph.add_run(text)
    _docx_set_run_font(run, bold=bold)
    return paragraph


def _docx_set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    _docx_style_paragraph(paragraph, after=0)
    run = paragraph.add_run(text.strip())
    _docx_set_run_font(run, bold=bold)


def _docx_add_markdown_table(doc, rows: list[str]) -> None:
    parsed = []
    for row in rows:
        cells = [cell.strip() for cell in row.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{2,}:?", cell.replace(" ", "")) for cell in cells):
            continue
        parsed.append(cells)
    if not parsed:
        return

    max_cols = max(len(row) for row in parsed)
    table = doc.add_table(rows=len(parsed), cols=max_cols)
    table.style = "Table Grid"
    table.autofit = True
    for r, row in enumerate(parsed):
        for c in range(max_cols):
            text = row[c] if c < len(row) else ""
            _docx_set_cell_text(table.cell(r, c), text, bold=(r == 0))
    doc.add_paragraph()


def _docx_flush_table(doc, table_rows: list[str]) -> None:
    if table_rows:
        _docx_add_markdown_table(doc, table_rows)


def export_docx(session: SessionState) -> dict[str, str]:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover
        raise RuntimeError("python-docx 未安装，无法导出 Word") from exc

    doc = Document()
    _docx_configure_document(doc)
    table_rows: list[str] = []

    def flush() -> None:
        nonlocal table_rows
        _docx_flush_table(doc, table_rows)
        table_rows = []

    for raw in session.prd_markdown.splitlines():
        line = raw.rstrip()
        if not line:
            flush()
            continue
        if line.startswith("|"):
            table_rows.append(line)
            continue
        flush()

        if line.startswith("# "):
            _docx_add_paragraph(doc, line[2:].strip(), style="Heading 1", bold=True)
        elif line.startswith("## "):
            _docx_add_paragraph(doc, line[3:].strip(), style="Heading 2", bold=True)
        elif line.startswith("### "):
            _docx_add_paragraph(doc, line[4:].strip(), style="Heading 3", bold=True)
        elif line.startswith("- "):
            _docx_add_paragraph(doc, line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            _docx_add_paragraph(doc, re.sub(r"^\d+\.\s+", "", line).strip(), style="List Number")
        else:
            _docx_add_paragraph(doc, line)
    flush()

    filename = f"{session.session_id}_{_safe_name(session.sub_type)}.docx"
    path = EXPORT_DIR / filename
    doc.save(path)
    return {"type": "docx", "filename": filename, "url": f"/exports/{filename}"}


def _pdf_escape(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _pdf_text_segments(text: str) -> list[tuple[str, str]]:
    segments: list[tuple[str, str]] = []
    current = ""
    current_kind = ""
    for ch in text:
        kind = "latin" if ord(ch) <= 127 else "cjk"
        if current and kind != current_kind:
            segments.append((current_kind, current))
            current = ch
        else:
            current += ch
        current_kind = kind
    if current:
        segments.append((current_kind, current))
    return segments


def _pdf_show_text_commands(text: str, bold: bool) -> list[str]:
    commands: list[str] = []
    for kind, segment in _pdf_text_segments(text):
        if kind == "latin":
            font = "/F3" if bold else "/F2"
            commands.append(f"{font} {FONT_SIZE_PT} Tf ({_pdf_escape(segment)}) Tj")
        else:
            hex_text = segment.encode("utf-16-be").hex().upper()
            commands.append(f"/F1 {FONT_SIZE_PT} Tf <{hex_text}> Tj")
    return commands


def _pdf_wrap(text: str, max_units: int = 92) -> list[str]:
    units = 0
    current = ""
    lines: list[str] = []
    for ch in text:
        width = 2 if ord(ch) > 127 else 1
        if units + width > max_units and current:
            lines.append(current)
            current = ch
            units = width
        else:
            current += ch
            units += width
    if current:
        lines.append(current)
    return lines or [""]


def _markdown_to_pdf_lines(markdown: str) -> list[tuple[str, bool]]:
    lines: list[tuple[str, bool]] = []
    in_table = False
    for raw in markdown.splitlines():
        line = raw.strip()
        if not line:
            in_table = False
            lines.append(("", False))
            continue
        if line.startswith("|"):
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            if all(re.fullmatch(r":?-{2,}:?", cell.replace(" ", "")) for cell in cells):
                continue
            text = " | ".join(cells)
            in_table = True
            for wrapped in _pdf_wrap(text, 86):
                lines.append((wrapped, False))
            continue
        in_table = False
        bold = False
        if line.startswith("# "):
            line = line[2:].strip()
            bold = True
        elif line.startswith("## "):
            line = line[3:].strip()
            bold = True
        elif line.startswith("### "):
            line = line[4:].strip()
            bold = True
        elif line.startswith("- "):
            line = f"• {line[2:].strip()}"
        for wrapped in _pdf_wrap(line, 92):
            lines.append((wrapped, bold))
    return lines


def _build_simple_pdf(lines: list[tuple[str, bool]], title: str) -> bytes:
    # PDF uses STSong-Light for CJK text and Helvetica/Helvetica-Bold for Latin fallback.
    page_width = 595
    page_height = 842
    left = 54
    top = 790
    line_height = 13
    bottom = 48
    objects: list[bytes] = []

    pages: list[list[tuple[str, bool]]] = []
    current: list[tuple[str, bool]] = []
    y = top
    for line in lines:
        if y < bottom:
            pages.append(current)
            current = []
            y = top
        current.append(line)
        y -= line_height
    if current:
        pages.append(current)

    def add(obj: str | bytes) -> int:
        if isinstance(obj, str):
            obj = obj.encode("utf-8")
        objects.append(obj)
        return len(objects)

    catalog_id = add(b"")
    pages_id = add(b"")
    font_cn_id = add(
        "<< /Type /Font /Subtype /Type0 /BaseFont /STSong-Light /Encoding /UniGB-UCS2-H "
        "/DescendantFonts [ << /Type /Font /Subtype /CIDFontType0 /BaseFont /STSong-Light "
        "/CIDSystemInfo << /Registry (Adobe) /Ordering (GB1) /Supplement 2 >> >> ] >>"
    )
    font_en_id = add("<< /Type /Font /Subtype /TrueType /BaseFont /Calibri /Encoding /WinAnsiEncoding >>")
    font_en_bold_id = add("<< /Type /Font /Subtype /TrueType /BaseFont /Calibri-Bold /Encoding /WinAnsiEncoding >>")

    page_ids: list[int] = []
    for page in pages:
        ascii_commands = ["BT", f"/F1 {FONT_SIZE_PT} Tf", f"{left} {top} Td", f"{line_height} TL"]
        for text, bold in page:
            if text:
                ascii_commands.extend(_pdf_show_text_commands(text, bold))
            ascii_commands.append("T*")
        ascii_commands.append("ET")
        content = "\n".join(ascii_commands).encode("ascii")
        content_id = add(b"<< /Length " + str(len(content)).encode("ascii") + b" >>\nstream\n" + content + b"\nendstream")
        page_id = add(
            f"<< /Type /Page /Parent {pages_id} 0 R /MediaBox [0 0 {page_width} {page_height}] "
            f"/Resources << /Font << /F1 {font_cn_id} 0 R /F2 {font_en_id} 0 R /F3 {font_en_bold_id} 0 R >> >> "
            f"/Contents {content_id} 0 R >>"
        )
        page_ids.append(page_id)

    objects[catalog_id - 1] = f"<< /Type /Catalog /Pages {pages_id} 0 R >>".encode("ascii")
    kids = " ".join(f"{page_id} 0 R" for page_id in page_ids)
    objects[pages_id - 1] = f"<< /Type /Pages /Kids [ {kids} ] /Count {len(page_ids)} >>".encode("ascii")

    output = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for idx, obj in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{idx} 0 obj\n".encode("ascii"))
        output.extend(obj)
        output.extend(b"\nendobj\n")
    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root {catalog_id} 0 R /Info << /Title ({_pdf_escape(title)}) >> >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n".encode("utf-8")
    )
    return bytes(output)


def export_pdf(session: SessionState) -> dict[str, str]:
    filename = f"{session.session_id}_{_safe_name(session.sub_type)}.pdf"
    path = EXPORT_DIR / filename
    lines = _markdown_to_pdf_lines(session.prd_markdown)
    title = session.sub_type or "PRD"
    path.write_bytes(_build_simple_pdf(lines, title))
    return {"type": "pdf", "filename": filename, "url": f"/exports/{filename}"}


def export_session(session: SessionState, export_type: str) -> dict[str, str]:
    if not session.prd_markdown:
        raise ValueError("当前会话还没有生成 PRD，无法导出")
    if export_type == "docx":
        return export_docx(session)
    if export_type == "pdf":
        return export_pdf(session)
    raise ValueError("仅支持 Word(.docx) 或 PDF(.pdf) 导出")
